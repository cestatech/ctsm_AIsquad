"""Generate Analysis Data Reviewer's Guide (ADRG) skeleton from ADaM content."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import HTTPException, status

_ADAM_DOCUMENT_TYPES = frozenset({"ADAM_DATASET", "ADAM_SPECIFICATION"})


def build_adrg_docx(content: dict) -> bytes:
    """
    Build an ADRG skeleton DOCX with one section per ADaM analysis dataset.

    Each section documents purpose, key variables, and derivation summary.
    """
    document_type = content.get("document_type", "")
    if document_type not in _ADAM_DOCUMENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={
                "code": "NOT_ADAM",
                "message": "Artifact content is not an ADaM dataset document.",
            },
        )

    datasets = content.get("datasets", [])
    if not datasets:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={
                "code": "NO_DATASETS",
                "message": "ADaM artifact has no datasets to export.",
            },
        )

    study_name = content.get("study_name") or content.get("protocol_number") or "Study"
    protocol = content.get("protocol_number") or study_name
    ig_version = content.get("adam_ig_version", "1.3")

    document = Document()
    title = document.add_heading("Analysis Data Reviewer's Guide (ADRG)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{study_name} — Protocol {protocol}")
    run.font.size = Pt(12)

    document.add_paragraph(
        f"CDISC ADaM IG {ig_version} — generated ADRG skeleton describing "
        "analysis dataset purpose, key variables, and derivation summaries."
    )

    for dataset in datasets:
        dataset_code = dataset.get("dataset", "UNK")
        document.add_heading(f"Dataset: {dataset_code}", level=1)

        purpose = dataset.get("label") or f"{dataset_code} analysis dataset"
        structure = dataset.get("structure", "")
        purpose_text = purpose
        if structure:
            purpose_text = f"{purpose} ({structure})"
        document.add_heading("Purpose", level=2)
        document.add_paragraph(purpose_text)

        key_vars = dataset.get("key_variables") or []
        if not key_vars:
            key_vars = [
                v.get("variable") or v.get("name")
                for v in dataset.get("variables", [])
                if isinstance(v, dict)
            ][:5]
        document.add_heading("Key Variables", level=2)
        if key_vars:
            for var_name in key_vars:
                document.add_paragraph(str(var_name), style="List Bullet")
        else:
            document.add_paragraph("No key variables specified.")

        document.add_heading("Derivation Summary", level=2)
        derivations = _collect_derivations(dataset)
        if derivations:
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            headers = table.rows[0].cells
            headers[0].text = "Variable"
            headers[1].text = "Label"
            headers[2].text = "Derivation"
            for var_name, label, derivation in derivations:
                row = table.add_row().cells
                row[0].text = var_name
                row[1].text = label
                row[2].text = derivation
        else:
            document.add_paragraph("No derivations documented.")

    trace_notes = content.get("traceability_notes", [])
    if trace_notes:
        document.add_heading("Traceability Notes", level=1)
        for note in trace_notes:
            document.add_paragraph(str(note), style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_adrg_filename(protocol_number: str | None, study_name: str | None) -> str:
    """Return a safe ADRG download filename."""
    base = protocol_number or study_name or "study"
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", str(base)).strip("_").lower() or "study"
    return f"adrg_{slug}.docx"


def build_adam_define_filename(protocol_number: str | None, study_name: str | None) -> str:
    """Return a safe ADaM define.xml download filename."""
    base = protocol_number or study_name or "study"
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", str(base)).strip("_").lower() or "study"
    return f"adam_{slug}_define.xml"


def _collect_derivations(dataset: dict) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for var in dataset.get("variables", []):
        if isinstance(var, str):
            continue
        name = var.get("variable") or var.get("name") or "UNK"
        label = var.get("label") or name
        derivation = (var.get("derivation") or "").strip()
        if derivation:
            rows.append((name, label, derivation))
    for flag in dataset.get("population_flags", []):
        name = flag.get("variable", "UNK")
        label = flag.get("label") or name
        derivation = (flag.get("derivation") or "").strip()
        if derivation:
            rows.append((name, label, derivation))
    return rows
