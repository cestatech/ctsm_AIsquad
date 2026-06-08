"""DOCX export for narrative clinical artifacts."""

from __future__ import annotations

import io
from datetime import UTC, datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SECTION_ORDER = {
    "PROTOCOL": [
        "synopsis",
        "objectives",
        "study_design",
        "eligibility",
        "endpoints",
        "statistical_considerations",
        "sections",
    ],
    "SAP": [
        "analysis_populations",
        "estimands",
        "primary_endpoint_analysis",
        "secondary_endpoints",
        "missing_data",
        "sections",
    ],
    "CSR": [
        "study_identification",
        "synopsis",
        "sections",
    ],
}


def _heading_text(key: str) -> str:
    return key.replace("_", " ").strip().title()


def _add_table(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    columns: list[str] = []
    for row in rows:
        if isinstance(row, dict):
            for key in row:
                if key not in columns:
                    columns.append(key)
    if not columns:
        return

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(columns):
        header_cells[index].text = _heading_text(column)
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows[:200]:
        if not isinstance(row, dict):
            continue
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row.get(column, ""))

    if len(rows) > 200:
        document.add_paragraph(f"Showing first 200 of {len(rows)} rows.")


def _add_value(document: Document, value: Any, level: int = 2) -> None:
    """Recursively append structured content to a Word document."""
    if value is None:
        document.add_paragraph("—")
        return

    if isinstance(value, str):
        for paragraph in value.split("\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        return

    if isinstance(value, (int, float, bool)):
        document.add_paragraph(str(value))
        return

    if isinstance(value, list):
        if not value:
            document.add_paragraph("(empty)")
            return
        if value and all(isinstance(item, dict) for item in value):
            _add_table(document, value)
            return
        for index, item in enumerate(value, start=1):
            if isinstance(item, dict):
                document.add_paragraph(f"Item {index}", style="List Number")
                _add_value(document, item, level + 1)
            else:
                document.add_paragraph(str(item), style="List Bullet")
        return

    if isinstance(value, dict):
        for key, nested in value.items():
            if nested in (None, "", [], {}):
                continue
            document.add_heading(_heading_text(str(key)), level=min(level, 3))
            _add_value(document, nested, level + 1)
        return

    document.add_paragraph(str(value))


def _add_title_page(
    document: Document,
    *,
    title: str,
    study_name: str,
    version_number: int,
    document_label: str,
) -> None:
    for _ in range(4):
        document.add_paragraph("")

    title_para = document.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    document.add_paragraph("")
    meta_lines = [
        f"Study: {study_name}",
        f"Document Type: {document_label}",
        f"Version: {version_number}",
        f"Exported: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}",
    ]
    for line in meta_lines:
        para = document.add_paragraph(line)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.color.rgb = RGBColor(71, 85, 105)

    document.add_page_break()


def _ordered_sections(content: dict, document_type: str) -> list[tuple[str, Any]]:
    order = SECTION_ORDER.get(document_type, [])
    seen: set[str] = set()
    sections: list[tuple[str, Any]] = []

    for key in order:
        value = content.get(key)
        if value not in (None, "", [], {}):
            sections.append((key, value))
            seen.add(key)

    skip_keys = {
        "document_type",
        "label",
        "synthetic_flag",
        "csv_files",
        "primary_csv_filename",
        "datasets",
        "ich_e3_compliant",
    }
    for key, value in content.items():
        if key in seen or key in skip_keys or value in (None, "", [], {}):
            continue
        sections.append((key, value))

    return sections


def export_docx(
    content: dict,
    *,
    title: str,
    study_name: str,
    version_number: int,
    document_label: str,
) -> bytes:
    """Render structured artifact JSON as a readable Word document."""
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    document_type = str(content.get("document_type") or document_label).upper()
    _add_title_page(
        document,
        title=title,
        study_name=study_name,
        version_number=version_number,
        document_label=document_label,
    )

    for key, value in _ordered_sections(content, document_type):
        document.add_heading(_heading_text(str(key)), level=1)
        _add_value(document, value, level=2)
        document.add_paragraph("")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
