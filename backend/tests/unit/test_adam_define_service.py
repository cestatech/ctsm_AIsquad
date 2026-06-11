"""Unit tests for ADaM define.xml builder."""

from __future__ import annotations

import io
import xml.etree.ElementTree as ET

import pytest
from docx import Document
from fastapi import HTTPException

from app.services.adam_define_service import build_adam_define_xml
from app.services.adrg_generation_service import build_adrg_docx

DEF_NS = "http://www.cdisc.org/ns/def/v2.1"


def _rich_adam_content() -> dict:
    return {
        "document_type": "ADAM_SPECIFICATION",
        "adam_ig_version": "1.3",
        "protocol_number": "STUDY-001",
        "study_name": "Example Study",
        "validation_engine": "internal",
        "datasets": [
            {
                "dataset": "ADSL",
                "label": "Subject Level Analysis Dataset",
                "structure": "One record per subject",
                "key_variables": ["STUDYID", "USUBJID"],
                "variables": [
                    {
                        "variable": "STUDYID",
                        "label": "Study Identifier",
                        "type": "Char",
                        "origin": "Assigned",
                        "derivation": "DM.STUDYID",
                    },
                    {
                        "variable": "USUBJID",
                        "label": "Unique Subject Identifier",
                        "type": "Char",
                        "origin": "SDTM.DM",
                        "derivation": "DM.USUBJID",
                    },
                    {
                        "variable": "AGE",
                        "label": "Age",
                        "type": "number",
                        "derivation": "Age at informed consent from RFICDTC and BRTHDTC",
                        "display_format": "3.0",
                    },
                ],
                "population_flags": [
                    {
                        "variable": "ITTFL",
                        "label": "Intent-to-Treat Flag",
                        "derivation": "Y if subject randomised",
                    }
                ],
            },
            {
                "dataset": "ADAE",
                "label": "Adverse Events Analysis Dataset",
                "structure": "One record per subject per event",
                "key_variables": ["USUBJID", "AEDECOD"],
                "variables": [
                    {
                        "variable": "AEDECOD",
                        "label": "Dictionary-Derived Term",
                        "type": "Char",
                        "derivation": "AE.AEDECOD",
                    }
                ],
                "population_flags": [],
            },
        ],
        "traceability_notes": ["ADSL derived from SDTM DM"],
    }


class TestBuildAdamDefineXml:
    def test_builds_xml_with_analysis_datasets(self):
        xml = build_adam_define_xml(_rich_adam_content())
        assert '<?xml' in xml
        assert "AnalysisDatasets" in xml
        assert "ADSL" in xml
        assert "STUDY-001" in xml
        assert "2.1" in xml

    def test_rejects_non_adam_content(self):
        with pytest.raises(HTTPException) as exc:
            build_adam_define_xml({"document_type": "SDTM_DATASET", "datasets": []})
        assert exc.value.status_code == 422

    def test_rejects_empty_datasets(self):
        with pytest.raises(HTTPException) as exc:
            build_adam_define_xml(
                {"document_type": "ADAM_SPECIFICATION", "datasets": []}
            )
        assert exc.value.status_code == 422


class TestAdamDefineXmlVariables:
    def test_includes_variable_metadata_and_display_format(self):
        xml = build_adam_define_xml(_rich_adam_content())
        assert 'Name="AGE"' in xml
        assert "Age" in xml
        assert "AnalysisVariableDisplayFormat" in xml
        assert "3.0" in xml
        assert 'Type="Derived"' in xml

    def test_methoddef_for_derived_variables(self):
        xml = build_adam_define_xml(_rich_adam_content())
        assert "MethodDef" in xml
        assert 'OID="MT.ADSL.AGE"' in xml
        assert "Age at informed consent" in xml

    def test_analysis_datasets_element_present(self):
        xml = build_adam_define_xml(_rich_adam_content())
        root = ET.fromstring(xml.split("\n", 1)[-1] if xml.startswith("<?xml") else xml)
        analysis = root.find(f".//{{{DEF_NS}}}AnalysisDatasets")
        assert analysis is not None
        datasets = analysis.findall(f"{{{DEF_NS}}}AnalysisDataset")
        assert len(datasets) == 2


class TestBuildAdrgDocx:
    def test_adrg_contains_expected_section_headings(self):
        body = build_adrg_docx(_rich_adam_content())
        document = Document(io.BytesIO(body))
        text_blocks = [p.text for p in document.paragraphs if p.text.strip()]
        assert "Analysis Data Reviewer's Guide (ADRG)" in text_blocks
        assert "Dataset: ADSL" in text_blocks
        assert "Dataset: ADAE" in text_blocks
        assert "Purpose" in text_blocks
        assert "Key Variables" in text_blocks
        assert "Derivation Summary" in text_blocks

    def test_adrg_includes_derivation_content(self):
        body = build_adrg_docx(_rich_adam_content())
        document = Document(io.BytesIO(body))
        text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        assert "DM.USUBJID" in text
        assert "Y if subject randomised" in text

    def test_rejects_non_adam_content(self):
        with pytest.raises(HTTPException) as exc:
            build_adrg_docx({"document_type": "PROTOCOL"})
        assert exc.value.status_code == 422
